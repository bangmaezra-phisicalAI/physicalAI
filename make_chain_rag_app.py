"""
Streamlit Make-Chain + RAG Chatbot Simulator

설치:
    pip install streamlit openai pypdf python-docx pandas numpy

실행:
    streamlit run make_chain_rag_app.py

특징:
- API 키를 웹페이지에서 직접 입력하며 파일이나 환경변수에 저장하지 않습니다.
- 에이전트의 이름/System Prompt와 일반 프롬프트 단계를 직접 편집할 수 있습니다.
- TXT, MD, PDF, DOCX, CSV 파일을 업로드하여 로컬 RAG 검색에 활용합니다.
- 최종 상호작용은 채팅 형태이며, 마지막 에이전트의 답변을 스트리밍합니다.
- 스캔 이미지형 PDF는 텍스트가 없으므로 별도의 OCR 처리가 필요합니다.
"""

from __future__ import annotations

import hashlib
import io
import re
from dataclasses import dataclass
from typing import Any, Callable

import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
from openai import OpenAI
from pypdf import PdfReader


# ---------------------------------------------------------------------
# 기본 설정
# ---------------------------------------------------------------------

st.set_page_config(
    page_title="Make Chain + RAG 챗봇",
    page_icon="🔗",
    layout="wide",
)

DEFAULT_SYSTEM_PROMPTS = [
    {
        "name": "자료 분석가",
        "prompt": (
            "당신은 업로드된 자료를 정확하게 분석하는 자료 분석가입니다.\n"
            "제공된 RAG 문맥을 우선 사용하세요.\n"
            "자료에 없는 사실은 지어내지 말고, 확인할 수 없다고 명시하세요.\n"
            "핵심 내용을 구조적으로 정리하세요."
        ),
    },
    {
        "name": "최종 답변 작성자",
        "prompt": (
            "당신은 친절하고 정확한 한국어 답변 작성자입니다.\n"
            "이전 단계의 결과와 제공된 RAG 문맥을 활용해 최종 답변을 작성하세요.\n"
            "근거가 부족하면 추측하지 마세요.\n"
            "가능하면 답변 끝에 참조한 파일명을 표시하세요."
        ),
    },
]

DEFAULT_CHAIN = [
    {"type": "prompt", "value": "다음 질문의 핵심 의도와 필요한 근거를 먼저 파악하세요."},
    {"type": "agent", "agent_index": 0},
    {"type": "prompt", "value": "위 분석을 바탕으로 사용자가 바로 이해할 수 있는 최종 답변을 작성하세요."},
    {"type": "agent", "agent_index": 1},
]


@dataclass
class DocumentChunk:
    text: str
    source: str
    chunk_id: int


# ---------------------------------------------------------------------
# 세션 상태
# ---------------------------------------------------------------------

def initialize_state() -> None:
    defaults: dict[str, Any] = {
        "agents": DEFAULT_SYSTEM_PROMPTS.copy(),
        "chain_steps": DEFAULT_CHAIN.copy(),
        "messages": [],
        "rag_signature": None,
        "rag_chunks": [],
        "rag_embeddings": None,
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


initialize_state()


# ---------------------------------------------------------------------
# 문서 추출 및 RAG
# ---------------------------------------------------------------------

def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text(uploaded_file: Any) -> str:
    """지원 파일에서 텍스트를 추출합니다."""
    suffix = uploaded_file.name.lower().rsplit(".", 1)[-1]
    raw = uploaded_file.getvalue()

    if suffix in {"txt", "md"}:
        for encoding in ("utf-8", "cp949", "euc-kr"):
            try:
                return normalize_text(raw.decode(encoding))
            except UnicodeDecodeError:
                continue
        return normalize_text(raw.decode("utf-8", errors="replace"))

    if suffix == "pdf":
        reader = PdfReader(io.BytesIO(raw))
        pages = []
        for page_number, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(f"[페이지 {page_number}]\n{page_text}")
        return normalize_text("\n\n".join(pages))

    if suffix == "docx":
        document = Document(io.BytesIO(raw))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                paragraphs.append(" | ".join(cells))
        return normalize_text("\n".join(paragraphs))

    if suffix == "csv":
        try:
            dataframe = pd.read_csv(io.BytesIO(raw))
        except UnicodeDecodeError:
            dataframe = pd.read_csv(io.BytesIO(raw), encoding="cp949")
        return normalize_text(dataframe.to_csv(index=False))

    raise ValueError(f"지원하지 않는 파일 형식입니다: {uploaded_file.name}")


def split_text(text: str, source: str, chunk_size: int, overlap: int) -> list[DocumentChunk]:
    """문단 경계를 최대한 유지하며 문자 단위로 청크를 만듭니다."""
    if not text:
        return []

    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[DocumentChunk] = []
    buffer = ""

    for paragraph in paragraphs:
        candidate = f"{buffer}\n\n{paragraph}".strip() if buffer else paragraph

        if len(candidate) <= chunk_size:
            buffer = candidate
            continue

        if buffer:
            chunks.append(
                DocumentChunk(text=buffer, source=source, chunk_id=len(chunks))
            )
            tail = buffer[-overlap:] if overlap > 0 else ""
            buffer = f"{tail}\n\n{paragraph}".strip()
        else:
            start = 0
            while start < len(paragraph):
                end = min(start + chunk_size, len(paragraph))
                chunks.append(
                    DocumentChunk(
                        text=paragraph[start:end],
                        source=source,
                        chunk_id=len(chunks),
                    )
                )
                if end == len(paragraph):
                    buffer = ""
                    break
                start = max(end - overlap, start + 1)

    if buffer:
        chunks.append(DocumentChunk(text=buffer, source=source, chunk_id=len(chunks)))

    return chunks


def file_signature(uploaded_files: list[Any], chunk_size: int, overlap: int) -> str:
    digest = hashlib.sha256()
    digest.update(str(chunk_size).encode())
    digest.update(str(overlap).encode())
    for file in uploaded_files:
        digest.update(file.name.encode("utf-8", errors="ignore"))
        digest.update(file.getvalue())
    return digest.hexdigest()


def batched(items: list[str], batch_size: int = 64):
    for start in range(0, len(items), batch_size):
        yield items[start : start + batch_size]


def create_embeddings(
    client: OpenAI,
    texts: list[str],
    embedding_model: str,
) -> np.ndarray:
    vectors: list[list[float]] = []
    for batch in batched(texts):
        response = client.embeddings.create(
            model=embedding_model,
            input=batch,
        )
        vectors.extend(item.embedding for item in response.data)

    matrix = np.asarray(vectors, dtype=np.float32)
    norms = np.linalg.norm(matrix, axis=1, keepdims=True)
    return matrix / np.clip(norms, 1e-12, None)


def retrieve_chunks(
    client: OpenAI,
    query: str,
    chunks: list[DocumentChunk],
    embeddings: np.ndarray,
    embedding_model: str,
    top_k: int,
    min_score: float,
) -> list[tuple[DocumentChunk, float]]:
    if not chunks or embeddings is None:
        return []

    response = client.embeddings.create(model=embedding_model, input=query)
    query_vector = np.asarray(response.data[0].embedding, dtype=np.float32)
    query_vector /= max(float(np.linalg.norm(query_vector)), 1e-12)

    scores = embeddings @ query_vector
    ranked_indices = np.argsort(scores)[::-1][:top_k]

    results = []
    for index in ranked_indices:
        score = float(scores[index])
        if score >= min_score:
            results.append((chunks[int(index)], score))
    return results


def format_rag_context(results: list[tuple[DocumentChunk, float]]) -> str:
    if not results:
        return (
            "[RAG 검색 결과 없음]\n"
            "업로드 자료에서 질문과 충분히 관련된 내용을 찾지 못했습니다."
        )

    blocks = []
    for number, (chunk, score) in enumerate(results, start=1):
        blocks.append(
            f"[근거 {number} | 파일: {chunk.source} | 유사도: {score:.3f}]\n"
            f"{chunk.text}"
        )
    return "\n\n".join(blocks)


# ---------------------------------------------------------------------
# LLM 에이전트와 체인
# ---------------------------------------------------------------------

def response(
    client: OpenAI,
    model: str,
    system: str,
    user: str,
    temperature: float,
    stream: bool = False,
):
    """사용자가 제시한 response 함수 구조를 유지한 Chat Completions 래퍼입니다."""
    return client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        temperature=temperature,
        stream=stream,
    )


def make_agent(
    client: OpenAI,
    model: str,
    system_prompt: str,
    temperature: float,
) -> Callable[[str, bool], Any]:
    def llm_agent(user: str, stream: bool = False):
        return response(
            client=client,
            model=model,
            system=system_prompt,
            user=user,
            temperature=temperature,
            stream=stream,
        )

    return llm_agent


def make_chain(*steps: Any) -> Callable[[str, bool], tuple[str, list[dict[str, str]]]]:
    """
    문자열은 현재 입력 앞에 붙이고, callable은 에이전트로 실행합니다.
    마지막 에이전트만 선택적으로 스트리밍합니다.
    """
    def chain(user_input: str, stream_final: bool = False):
        current = user_input
        trace: list[dict[str, str]] = []
        callable_positions = [i for i, step in enumerate(steps) if callable(step)]
        final_callable_position = callable_positions[-1] if callable_positions else -1

        for index, step in enumerate(steps):
            if callable(step):
                use_stream = stream_final and index == final_callable_position
                result = step(current, stream=use_stream)

                if use_stream:
                    return result, trace

                current = result.choices[0].message.content or ""
                trace.append(
                    {
                        "step": f"에이전트 단계 {index + 1}",
                        "output": current,
                    }
                )

            elif isinstance(step, str):
                current = f"{step}\n\n{current}".strip()
                trace.append(
                    {
                        "step": f"프롬프트 단계 {index + 1}",
                        "output": current,
                    }
                )

        return current, trace

    return chain


def build_runtime_steps(
    client: OpenAI,
    model: str,
    agents: list[dict[str, str]],
    chain_steps: list[dict[str, Any]],
    temperature: float,
) -> list[Any]:
    runtime_steps: list[Any] = []

    for step in chain_steps:
        if step["type"] == "prompt":
            runtime_steps.append(step.get("value", ""))
        elif step["type"] == "agent":
            agent_index = int(step.get("agent_index", 0))
            if not 0 <= agent_index < len(agents):
                raise ValueError(f"존재하지 않는 에이전트 번호입니다: {agent_index + 1}")
            runtime_steps.append(
                make_agent(
                    client=client,
                    model=model,
                    system_prompt=agents[agent_index]["prompt"],
                    temperature=temperature,
                )
            )

    return runtime_steps


# ---------------------------------------------------------------------
# 사이드바: API와 RAG 설정
# ---------------------------------------------------------------------

with st.sidebar:
    st.header("⚙️ 실행 설정")

    api_key = st.text_input(
        "OpenAI API Key",
        type="password",
        placeholder="sk-...",
        help="입력값은 현재 Streamlit 세션에서만 사용하며 앱이 파일에 저장하지 않습니다.",
    )

    model = st.text_input(
        "GPT 모델",
        value="gpt-5",
        help="본인 API 계정에서 사용 가능한 모델명을 입력하세요.",
    )

    embedding_model = st.text_input(
        "임베딩 모델",
        value="text-embedding-3-small",
    )

    temperature = st.slider(
        "Temperature",
        min_value=0.0,
        max_value=1.5,
        value=0.2,
        step=0.1,
    )

    st.divider()
    st.subheader("📚 RAG 설정")

    uploaded_files = st.file_uploader(
        "참조 파일 업로드",
        type=["txt", "md", "pdf", "docx", "csv"],
        accept_multiple_files=True,
    )

    chunk_size = st.slider("청크 크기", 300, 3000, 1000, 100)
    overlap = st.slider(
        "청크 겹침",
        0,
        min(500, chunk_size - 1),
        min(150, chunk_size - 1),
        25,
    )
    top_k = st.slider("검색할 청크 수", 1, 10, 4)
    min_score = st.slider(
        "최소 유사도",
        min_value=0.0,
        max_value=1.0,
        value=0.25,
        step=0.05,
        help="너무 높으면 검색 결과가 없어질 수 있습니다.",
    )

    strict_rag = st.checkbox(
        "자료에 근거가 없으면 답변 제한",
        value=True,
    )

    if st.button("대화 기록 초기화", use_container_width=True):
        st.session_state.messages = []
        st.rerun()


# ---------------------------------------------------------------------
# 메인 화면: 편집기
# ---------------------------------------------------------------------

st.title("🔗 Make Chain + RAG 챗봇 시뮬레이터")
st.caption(
    "에이전트와 문자열 프롬프트를 원하는 순서로 연결하고, "
    "업로드 파일을 검색해 답변에 참조시키는 실습용 앱입니다."
)

editor_tab, rag_tab, chat_tab = st.tabs(
    ["🧠 에이전트·체인 편집", "📚 RAG 상태", "💬 챗봇"]
)

with editor_tab:
    st.subheader("에이전트 편집")

    remove_agent_index = None
    for index, agent in enumerate(st.session_state.agents):
        with st.expander(
            f"에이전트 {index + 1}: {agent['name']}",
            expanded=index == 0,
        ):
            agent["name"] = st.text_input(
                "에이전트 이름",
                value=agent["name"],
                key=f"agent_name_{index}",
            )
            agent["prompt"] = st.text_area(
                "System Prompt",
                value=agent["prompt"],
                height=180,
                key=f"agent_prompt_{index}",
            )
            if st.button(
                "이 에이전트 삭제",
                key=f"remove_agent_{index}",
                disabled=len(st.session_state.agents) <= 1,
            ):
                remove_agent_index = index

    if remove_agent_index is not None:
        st.session_state.agents.pop(remove_agent_index)
        for step in st.session_state.chain_steps:
            if step["type"] == "agent":
                step["agent_index"] = min(
                    int(step.get("agent_index", 0)),
                    len(st.session_state.agents) - 1,
                )
        st.rerun()

    if st.button("➕ 에이전트 추가"):
        st.session_state.agents.append(
            {
                "name": f"새 에이전트 {len(st.session_state.agents) + 1}",
                "prompt": "당신의 역할과 답변 규칙을 여기에 입력하세요.",
            }
        )
        st.rerun()

    st.divider()
    st.subheader("체인 단계 편집")
    st.info(
        "문자열 프롬프트는 현재 결과 앞에 추가되고, 에이전트 단계는 그 내용을 LLM에 전달합니다."
    )

    remove_step_index = None
    move_up_index = None
    move_down_index = None

    agent_labels = [
        f"{i + 1}. {agent['name']}"
        for i, agent in enumerate(st.session_state.agents)
    ]

    for index, step in enumerate(st.session_state.chain_steps):
        with st.container(border=True):
            left, middle, right = st.columns([5, 1, 1])

            with left:
                step_type = st.selectbox(
                    f"단계 {index + 1} 종류",
                    options=["prompt", "agent"],
                    index=0 if step["type"] == "prompt" else 1,
                    format_func=lambda x: "문자열 프롬프트" if x == "prompt" else "LLM 에이전트",
                    key=f"step_type_{index}",
                )

                if step_type != step["type"]:
                    if step_type == "prompt":
                        st.session_state.chain_steps[index] = {
                            "type": "prompt",
                            "value": "새 프롬프트를 입력하세요.",
                        }
                    else:
                        st.session_state.chain_steps[index] = {
                            "type": "agent",
                            "agent_index": 0,
                        }
                    st.rerun()

                if step["type"] == "prompt":
                    step["value"] = st.text_area(
                        "프롬프트",
                        value=step.get("value", ""),
                        height=100,
                        key=f"step_value_{index}",
                    )
                else:
                    selected_index = min(
                        int(step.get("agent_index", 0)),
                        len(agent_labels) - 1,
                    )
                    step["agent_index"] = st.selectbox(
                        "실행할 에이전트",
                        options=list(range(len(agent_labels))),
                        index=selected_index,
                        format_func=lambda i: agent_labels[i],
                        key=f"step_agent_{index}",
                    )

            with middle:
                st.write("")
                st.write("")
                if st.button("⬆️", key=f"up_{index}", disabled=index == 0):
                    move_up_index = index
                if st.button(
                    "⬇️",
                    key=f"down_{index}",
                    disabled=index == len(st.session_state.chain_steps) - 1,
                ):
                    move_down_index = index

            with right:
                st.write("")
                st.write("")
                if st.button(
                    "🗑️",
                    key=f"delete_step_{index}",
                    disabled=len(st.session_state.chain_steps) <= 1,
                ):
                    remove_step_index = index

    if remove_step_index is not None:
        st.session_state.chain_steps.pop(remove_step_index)
        st.rerun()

    if move_up_index is not None:
        steps = st.session_state.chain_steps
        steps[move_up_index - 1], steps[move_up_index] = (
            steps[move_up_index],
            steps[move_up_index - 1],
        )
        st.rerun()

    if move_down_index is not None:
        steps = st.session_state.chain_steps
        steps[move_down_index + 1], steps[move_down_index] = (
            steps[move_down_index],
            steps[move_down_index + 1],
        )
        st.rerun()

    add_left, add_right = st.columns(2)
    with add_left:
        if st.button("➕ 문자열 프롬프트 단계 추가", use_container_width=True):
            st.session_state.chain_steps.append(
                {"type": "prompt", "value": "새 프롬프트를 입력하세요."}
            )
            st.rerun()
    with add_right:
        if st.button("➕ 에이전트 단계 추가", use_container_width=True):
            st.session_state.chain_steps.append(
                {"type": "agent", "agent_index": 0}
            )
            st.rerun()


# ---------------------------------------------------------------------
# RAG 인덱스 생성
# ---------------------------------------------------------------------

rag_error = None
rag_notice = None

if uploaded_files and api_key:
    try:
        signature = file_signature(uploaded_files, chunk_size, overlap)

        if signature != st.session_state.rag_signature:
            client_for_index = OpenAI(api_key=api_key)
            chunks: list[DocumentChunk] = []
            empty_files = []

            with st.spinner("업로드 파일을 읽고 RAG 인덱스를 생성하는 중입니다..."):
                for uploaded_file in uploaded_files:
                    extracted = extract_text(uploaded_file)
                    if not extracted.strip():
                        empty_files.append(uploaded_file.name)
                        continue
                    chunks.extend(
                        split_text(
                            text=extracted,
                            source=uploaded_file.name,
                            chunk_size=chunk_size,
                            overlap=overlap,
                        )
                    )

                if chunks:
                    embeddings = create_embeddings(
                        client=client_for_index,
                        texts=[chunk.text for chunk in chunks],
                        embedding_model=embedding_model,
                    )
                else:
                    embeddings = None

            st.session_state.rag_signature = signature
            st.session_state.rag_chunks = chunks
            st.session_state.rag_embeddings = embeddings

            if empty_files:
                rag_notice = (
                    "텍스트를 추출하지 못한 파일: "
                    + ", ".join(empty_files)
                    + ". 스캔 PDF라면 OCR이 필요합니다."
                )

    except Exception as error:
        rag_error = str(error)

elif not uploaded_files:
    st.session_state.rag_signature = None
    st.session_state.rag_chunks = []
    st.session_state.rag_embeddings = None


with rag_tab:
    st.subheader("RAG 인덱스 상태")

    if rag_error:
        st.error(f"RAG 인덱스 생성 오류: {rag_error}")
    elif uploaded_files and not api_key:
        st.warning("임베딩을 생성하려면 사이드바에 API 키를 입력하세요.")
    elif st.session_state.rag_chunks:
        source_counts: dict[str, int] = {}
        for chunk in st.session_state.rag_chunks:
            source_counts[chunk.source] = source_counts.get(chunk.source, 0) + 1

        st.success(
            f"{len(source_counts)}개 파일에서 "
            f"{len(st.session_state.rag_chunks)}개 청크를 생성했습니다."
        )
        st.dataframe(
            pd.DataFrame(
                [
                    {"파일명": source, "청크 수": count}
                    for source, count in source_counts.items()
                ]
            ),
            hide_index=True,
            use_container_width=True,
        )

        with st.expander("청크 미리보기"):
            for chunk in st.session_state.rag_chunks[:10]:
                st.markdown(f"**{chunk.source} · 청크 {chunk.chunk_id + 1}**")
                st.text(chunk.text[:1000])
    else:
        st.info("사이드바에서 참조 파일을 업로드하면 여기에 RAG 상태가 표시됩니다.")

    if rag_notice:
        st.warning(rag_notice)

    st.caption(
        "현재 버전은 파일 텍스트를 앱 메모리에만 보관합니다. "
        "페이지를 새로 열거나 서버가 재시작되면 다시 인덱싱됩니다."
    )


# ---------------------------------------------------------------------
# 챗봇
# ---------------------------------------------------------------------

with chat_tab:
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
            if message.get("sources"):
                with st.expander("참조 근거"):
                    for source in message["sources"]:
                        st.markdown(
                            f"**{source['file']}** · 유사도 {source['score']:.3f}"
                        )
                        st.text(source["text"])

    user_input = st.chat_input("질문을 입력하세요")

    if user_input:
        st.session_state.messages.append(
            {"role": "user", "content": user_input}
        )
        with st.chat_message("user"):
            st.markdown(user_input)

        with st.chat_message("assistant"):
            if not api_key:
                error_message = "사이드바에 OpenAI API 키를 입력해 주세요."
                st.error(error_message)
                st.session_state.messages.append(
                    {"role": "assistant", "content": error_message}
                )
                st.stop()

            if not any(
                step["type"] == "agent"
                for step in st.session_state.chain_steps
            ):
                error_message = "체인에 LLM 에이전트 단계를 하나 이상 추가해 주세요."
                st.error(error_message)
                st.session_state.messages.append(
                    {"role": "assistant", "content": error_message}
                )
                st.stop()

            try:
                client = OpenAI(api_key=api_key)

                retrieved = []
                if (
                    st.session_state.rag_chunks
                    and st.session_state.rag_embeddings is not None
                ):
                    retrieved = retrieve_chunks(
                        client=client,
                        query=user_input,
                        chunks=st.session_state.rag_chunks,
                        embeddings=st.session_state.rag_embeddings,
                        embedding_model=embedding_model,
                        top_k=top_k,
                        min_score=min_score,
                    )

                if strict_rag and uploaded_files and not retrieved:
                    final_text = (
                        "업로드된 자료에서 질문에 답할 만큼 관련된 근거를 "
                        "찾지 못했습니다. 질문을 더 구체적으로 작성하거나 "
                        "RAG 최소 유사도 값을 낮춰 보세요."
                    )
                    st.markdown(final_text)
                    st.session_state.messages.append(
                        {"role": "assistant", "content": final_text}
                    )
                    st.stop()

                rag_context = format_rag_context(retrieved)

                history = st.session_state.messages[-8:-1]
                history_text = "\n".join(
                    f"{'사용자' if item['role'] == 'user' else '도우미'}: "
                    f"{item['content']}"
                    for item in history
                )

                augmented_input = f"""
[대화 기록]
{history_text or "이전 대화 없음"}

[RAG 참조 문맥]
{rag_context}

[현재 사용자 질문]
{user_input}

[답변 규칙]
- RAG 근거가 제공된 경우 해당 근거를 우선 사용하세요.
- 근거에 없는 세부 사실은 추측하지 마세요.
- 최종 답변은 사용자의 질문에 직접 답하세요.
""".strip()

                runtime_steps = build_runtime_steps(
                    client=client,
                    model=model,
                    agents=st.session_state.agents,
                    chain_steps=st.session_state.chain_steps,
                    temperature=temperature,
                )
                chain = make_chain(*runtime_steps)

                stream_result, trace = chain(
                    augmented_input,
                    stream_final=True,
                )

                # 마지막 단계가 에이전트이면 OpenAI 스트림 객체가 반환됩니다.
                if hasattr(stream_result, "__iter__") and not isinstance(
                    stream_result, str
                ):
                    def token_generator():
                        for chunk in stream_result:
                            content = chunk.choices[0].delta.content or ""
                            if content:
                                yield content

                    final_text = st.write_stream(token_generator())
                else:
                    final_text = str(stream_result)
                    st.markdown(final_text)

                sources = [
                    {
                        "file": chunk.source,
                        "score": score,
                        "text": chunk.text,
                    }
                    for chunk, score in retrieved
                ]

                if sources:
                    with st.expander("참조 근거"):
                        for source in sources:
                            st.markdown(
                                f"**{source['file']}** · "
                                f"유사도 {source['score']:.3f}"
                            )
                            st.text(source["text"])

                if trace:
                    with st.expander("체인 중간 결과"):
                        for item in trace:
                            st.markdown(f"**{item['step']}**")
                            st.text(item["output"])

                st.session_state.messages.append(
                    {
                        "role": "assistant",
                        "content": final_text,
                        "sources": sources,
                    }
                )

            except Exception as error:
                error_message = f"실행 중 오류가 발생했습니다: {error}"
                st.error(error_message)
                st.session_state.messages.append(
                    {"role": "assistant", "content": error_message}
                )
